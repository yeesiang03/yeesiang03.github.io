from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "resume"
DOCX_PATH = OUTPUT / "Ku_Yee_Siang_Resume.docx"

FONT = "Aptos"
INK = RGBColor(26, 32, 44)
ACCENT = RGBColor(27, 92, 135)
MUTED = RGBColor(81, 94, 107)
RULE = "B9C9D6"


def set_font(run, size=9.2, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_bottom_border(paragraph, color=RULE, size="8", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=10.3, bold=True, color=ACCENT)
    run.font.letter_spacing = Pt(0.3)
    add_bottom_border(p)
    return p


def add_project(doc, title, label, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, size=9.5, bold=True)
    r = p.add_run(f"  |  {label}")
    set_font(r, size=8.6, italic=True, color=MUTED)

    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.13)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.2)
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            run.clear()
        run = p.add_run(bullet)
        set_font(run, size=8.65)


def build():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.46)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.56)
    section.right_margin = Inches(0.56)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    bullet_style = styles["List Bullet"]
    bullet_style.font.name = FONT
    bullet_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet_style.font.size = Pt(8.65)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("KU YEE SIANG")
    set_font(r, size=22, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(2)
    r = subtitle.add_run("AI AUTOMATION & AGENT ENGINEER")
    set_font(r, size=10.4, bold=True, color=MUTED)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(4)
    r = contact.add_run(
        "+60 17-421 6727  |  yeesiangku@gmail.com\n"
        "github.com/yeesiang03  |  linkedin.com/in/yee-siang-ku-11b069315"
    )
    set_font(r, size=8.4, color=MUTED)
    add_bottom_border(contact, color="6F8FA7", size="10", space="4")

    add_section_heading(doc, "Professional Profile")
    profile = doc.add_paragraph()
    profile.paragraph_format.space_after = Pt(1)
    profile.paragraph_format.line_spacing = 1.03
    r = profile.add_run(
        "Artificial Intelligence graduate focused on AI automation, agentic systems, and end-to-end product delivery. "
        "Built secure multi-agent workflows, LLM evaluation benchmarks, retrieval-augmented applications, and "
        "full-stack business tools using Python and TypeScript. Skilled at translating requirements into working "
        "solutions, integrating APIs and databases, and improving reliability through validation and safety guardrails."
    )
    set_font(r, size=8.9)

    add_section_heading(doc, "Education & Award")
    edu = doc.add_paragraph()
    edu.paragraph_format.space_after = Pt(0)
    r = edu.add_run("Beijing Institute of Technology")
    set_font(r, size=9.4, bold=True)
    r = edu.add_run("  |  Bachelor of Science in Artificial Intelligence  |  2022-2026")
    set_font(r, size=8.8)
    edu2 = doc.add_paragraph()
    edu2.paragraph_format.space_after = Pt(0)
    r = edu2.add_run("GPA: 3.6/4.0  |  Best Student Award")
    set_font(r, size=8.8, bold=True, color=ACCENT)

    add_section_heading(doc, "Selected Projects")
    add_project(
        doc,
        "Harness Agent - Enterprise Procurement Workbench",
        "Agent Engineering",
        [
            "Built an AI-powered procurement workbench that automates product sourcing, supplier comparison, inventory analysis, purchase-order workflows, product onboarding, and report generation.",
            "Engineered the Vue 3 and FastAPI platform with DeepSeek, MCP/ERP tools, specialist sub-agents, SSE execution streaming, human approval, long-term memory, progressive skills, and per-user OpenSandbox isolation.",
        ],
    )
    add_project(
        doc,
        "Medical LLM Causal Consistency Evaluation Benchmark",
        "Graduation Thesis",
        [
            "Designed and built MedCausalBenchmark from six NCCN clinical guidelines using a dual-axis taxonomy spanning medical scenarios and Basic, Joint, Nested, and Conditional counterfactual reasoning.",
            "Created a repeatable evaluation workflow comparing Pure LLM, RAG, and CausalSCM across Qwen, LLaMA, and DeepSeek; CausalSCM achieved 91.39% average accuracy.",
        ],
    )
    add_project(
        doc,
        "Mini-Map - Gemini-Powered Virtual Travel Journal",
        "Google Cloud Rapid Agent Hackathon",
        [
            "Developed a choice-driven travel application with Gemini Pro, Google Cloud Agent Builder, FastAPI, Next.js, and MongoDB Atlas, turning user preferences into progressive journeys constrained by geography and budget.",
            "Integrated validated MCP tools, GeoJSON persistence, Atlas Vector Search with 1024-dimensional embeddings, external API grounding, deduplication safeguards, and one-click itinerary export.",
        ],
    )
    add_project(
        doc,
        "Agentic Medical Chat System",
        "Full-Stack AI",
        [
            "Built an end-to-end medical Q&A application combining tool-augmented LLM agents, retrieval, a web interface, and safety guardrails for structured clinical conversations.",
            "Integrated frontend, agent orchestration, backend retrieval, validation, and error handling into a traceable full-stack workflow for safer medical information delivery.",
        ],
    )

    add_section_heading(doc, "Technical Skills")
    skills = [
        ("Programming", "Python, TypeScript, JavaScript, SQL, HTML/CSS"),
        ("Agentic AI & LLM", "Agent orchestration, multi-agent systems, LLM evaluation, benchmark design, RAG, prompt and context engineering, tool calling, human-in-the-loop, memory, safety guardrails"),
        ("Models & Platforms", "DeepSeek, Gemini Pro, Qwen, LLaMA, OpenAI Codex, Google Cloud Agent Builder, Vertex AI, MCP, OpenSandbox"),
        ("Full-Stack", "FastAPI, React, Next.js, Vue 3, Tailwind CSS, REST APIs, Server-Sent Events, API integration"),
        ("Data & Search", "MongoDB Atlas, Atlas Search, Vector Search, SQLite, embeddings, GeoJSON"),
        ("Engineering Practices", "Requirements analysis, debugging, testing, validation, access control, data isolation, Git/GitHub, Vercel"),
    ]
    table = doc.add_table(rows=len(skills), cols=2)
    set_table_geometry(table, [1450, 9185])
    for row, (label, value) in zip(table.rows, skills):
        row.cells[0].vertical_alignment = 1
        row.cells[1].vertical_alignment = 1
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0.5)
        r = p.add_run(label)
        set_font(r, size=8.35, bold=True, color=ACCENT)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0.5)
        r = p.add_run(value)
        set_font(r, size=8.25)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Ku Yee Siang | AI Automation & Agent Engineer")
    set_font(r, size=7.5, color=MUTED)

    props = doc.core_properties
    props.title = "Ku Yee Siang - AI Automation & Agent Engineer Resume"
    props.subject = "AI Automation & Agent Engineer Resume"
    props.author = "Ku Yee Siang"
    props.keywords = "AI Automation, Agent Engineer, LLM Evaluation, Agentic AI, FastAPI, Python, TypeScript"

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
